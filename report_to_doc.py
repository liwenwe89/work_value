# -*- coding: utf-8 -*-
import easygui
import xlrd
import docx
from docx.oxml.ns import qn
from docx.shared import Pt
import time

#应该不会超过99所以简单一些
def NumberToChinese(num):
    num_dict={"0":u"","1":u"一","2":u"二","3":u"三","4":u"四","5":u"五","6":u"六","7":u"七","8":u"八","9":u"九","10":u"十"}
 
    zh_num_str =""
    if(num>=10 and num < 20):
        zh_num_str = "十" + num_dict[str(num%10)]
    elif (num<10):
        zh_num_str =  num_dict[str(num)]
    else:
        zh_num_str = num_dict[str(int(num/10))]+ num_dict[str(num%10)]
    return zh_num_str


def open_worksheet(path):
    
    print(path)

    excel_book = xlrd.open_workbook(path) 

    #获取第一个sheet页面
    excel_table_1 = excel_book.sheet_by_index(0)
    print(excel_table_1.name)

    #获取有效行数
    sheet1_nrows = excel_table_1.nrows
    sheet1_ncols = excel_table_1.ncols

    #第一行为基本信息
    thisweekprocess_col = 0 # 本周进展
    nextweekplan_col = 0 #下周计划
    riskandprob_col = 0 #风险和问题
    initplan_col = 0 #初始计划
    status_col = 0 #当前状态
    report_col = 0 #是否汇报
    proname_col = 0 #项目名称
    productline_col = 0 # 产品线

    list_val = [[[]for i in range(sheet1_ncols)] for i in range(sheet1_nrows)]
    


    for i in range(0,sheet1_ncols):
        list_val[0][i] = excel_table_1.cell(0,i).value
        #    print(list_val[1][i] )
        if(r"初始计划" in list_val[0][i] ):
            initplan_col = i # 从0开始 所以减1 
        if(r"当前状态" in list_val[0][i]):
            status_col = i 
        if(r"汇报" in list_val[0][i]):
            report_col = i 
        if(r"项目名称" in list_val[0][i]):
            proname_col = i 
        if(r"产品线" in list_val[0][i]):
            productline_col = i 
        if(r"本周进展" in list_val[0][i]):
            thisweekprocess_col = i 
        else:
            thisweekprocess_col = initplan_col -3
        if(r"下周计划" in list_val[0][i]):
            nextweekplan_col = i 
        else:
            nextweekplan_col = initplan_col -2
        if(r"风险" in list_val[0][i]):
            riskandprob_col = i 
        else:
            riskandprob_col = initplan_col -1

    print(initplan_col , #初始计划
    status_col , #状态栏
    report_col #是否汇报
    )
    temp =""
    for i in range(1,sheet1_nrows): 
        #出现\t和\r重复重现的情况t
        #考虑使用正则表达式去除？？
        for j in (productline_col,proname_col,
                    initplan_col-3,initplan_col-2,initplan_col-1, # 进度/计划/风险 在初始计划的前几列
                    initplan_col,status_col,report_col):
            temp = str(excel_table_1.cell(i,j).value)
            if('\n' in temp and '\r' in temp):
                list_val[i][j] = temp.replace('\r','')
            elif ('\r' in temp):
                list_val[i][j] = temp.replace('\r','\n')
            else:
                list_val[i][j] = temp   
    dict_col ={"产品线":productline_col,"项目名称":proname_col,
                "本周进展":thisweekprocess_col,"风险":riskandprob_col,
                "状态":status_col,"汇报":report_col}
    return [dict_col,list_val]


def write_docx(read_excel_result,productline,file_docx):
    doct_col =  read_excel_result[0]
    list_val = read_excel_result[1]
    #处理行业产品线数据，由于段落选择没有合适的命令，所以区分行业，渠道，海外进行处理
    p_hy=file_docx.add_paragraph('')
    p_hy1 = p_hy.add_run(productline+'产品线')
    p_hy1.font.bold = True
    p_hy1.font.underline = True
    #颜色
    p_hy1.font.color.rgb = docx.shared.RGBColor(0, 176 , 80)
    # 设置中文字体
    p_hy1.font.name = u'微软雅黑'
    p_hy1.element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
    #字号
    p_hy1.font.size = Pt(11)
    pro_rep_num = 0

    for i in range(1,len(list_val)): 
        if( (productline in list_val[i][doct_col["产品线"]]) and list_val[i][doct_col["状态"]] == "开发中" and list_val[i][doct_col["汇报"]] == "是") :
            pro_rep_num += 1 
            temp = NumberToChinese(pro_rep_num)+"、"+ str(list_val[i][doct_col["项目名称"]]).replace("\r",'').replace('\n','')
            file_docx.add_paragraph(temp)

            temp = list_val[i][doct_col["本周进展"]]
            if(temp == ""):
                    print("error!") #考虑补充错误补充
            file_docx.add_paragraph(temp)
            p_hy = file_docx.add_paragraph("")
            temp_run = p_hy.add_run("风险以及问题：")
            temp_run.font.color.rgb = docx.shared.RGBColor(255,0,0)
            p_hy = file_docx.add_paragraph("")
            temp = list_val[i][doct_col["风险"]]
            if(temp == ""):
                temp = "1、暂无;"
            temp_run = p_hy.add_run(temp)
            temp_run.font.color.rgb = docx.shared.RGBColor(255,0,0)
            p_hy.add_run("\n")


if __name__ == '__main__':

    path = easygui.fileopenbox()
    read_excel_result = open_worksheet(path)

    file_docx=docx.Document()
    #file_docx.add_heading('行业组重点项目进展', 1)
    docx_style = file_docx.styles['Normal']
    # 设置西文字体
    docx_style.font.name = '微软雅黑'
    # 设置中文字体
    docx_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    write_docx(read_excel_result,"行业",file_docx)
    write_docx(read_excel_result,"渠道",file_docx)
    write_docx(read_excel_result,"专用",file_docx)
    write_docx(read_excel_result,"海外",file_docx)

    date_to = time.localtime()
    week_to = time.strftime("%U",date_to)
    file_docx.save("本周重点问题汇总" +"_W"+ week_to +".docx")

    input("已经完成\n"+"本周重点问题汇总" +"_W"+ week_to +".docx")