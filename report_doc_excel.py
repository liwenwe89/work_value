# -*- coding: utf-8 -*-
import easygui
import xlrd
import xlwt
import docx
from docx.oxml.ns import qn
from docx.shared import Pt
import time

#应该不会超过99所以简单一些
def NumberToChinese(num):
    num_dict={0:"","1":u"一","2":u"二","3":u"三","4":u"四","5":u"五","6":u"六","7":u"七","8":u"八","9":u"九","10":u"十"}
 
    zh_num_str =""
    
    if(num>=10 and num < 20):
        zh_num_str = "十" + num_dict[str(num%10)]
    elif (num<10):
        zh_num_str =  num_dict[str(num)]
    else:
        zh_num_str = num_dict[str(int(num/10))]+"十"+num_dict[str(num%10)]
    return zh_num_str


def open_worksheet(path):
    
    print(path)

    excel_book = xlrd.open_workbook(path) 

    #获取第一个sheet页面
    excel_table_1 = excel_book.sheet_by_index(0)
    print(excel_table_1.name)

    #获取有效行数
    sheet1_nrows = excel_table_1.nrows
    sheet1_ncols = excel_table_1.ncols

    #第一行为基本信息
    thisweekprocess_col = 0 # 本周进展
    nextweekplan_col = 0 #下周计划
    riskandprob_col = 0 #风险和问题
    initplan_col = 0 #初始计划
    status_col = 0 #当前状态
    report_col = 0 #是否汇报
    proname_col = 0 #项目名称
    productline_col = 0 # 产品线

    list_val = [[[]for i in range(sheet1_ncols)] for i in range(sheet1_nrows)]
    


    for i in range(0,sheet1_ncols):
        list_val[0][i] = excel_table_1.cell(0,i).value
        #    print(list_val[1][i] )
        
        if(r"项目经理" in list_val[0][i] ):
            pm_name_clo = i 
        
        if(r"核心" in list_val[0][i] ):
            corerequire_col = i  

        if(r"调整后计划" in list_val[0][i] ):
            adustplan_col = i  
        
        if(r"偏差" in list_val[0][i] ):
            deviation_col = i  

        if(r"初始计划" in list_val[0][i] ):
            initplan_col = i # 从0开始 所以减1 
        if(r"状态" in list_val[0][i]):
            status_col = i 
        if(r"汇报" in list_val[0][i]):
            report_col = i 
        if(r"项目名称" in list_val[0][i]):
            proname_col = i 
        if(r"产品线" in list_val[0][i]):
            productline_col = i 
        if(r"本周进展" in list_val[0][i]):
            thisweekprocess_col = i 
        else:
            thisweekprocess_col = initplan_col -3
        if(r"下周计划" in list_val[0][i]):
            nextweekplan_col = i 
        else:
            nextweekplan_col = initplan_col -2
        if(r"风险" in list_val[0][i]):
            riskandprob_col = i 
        else:
            riskandprob_col = initplan_col -1

    print(initplan_col , #初始计划
    status_col , #状态栏
    report_col #是否汇报
    )
    temp =""
    for i in range(1,sheet1_nrows): 
        #出现\t和\r重复重现的情况t
        #考虑使用正则表达式去除？？
    #    for j in (productline_col,proname_col,
    #                initplan_col-3,initplan_col-2,initplan_col-1, # 进度/计划/风险 在初始计划的前几列
    #                initplan_col,status_col,report_col):
    
        for j in range(0,sheet1_ncols):#range不包括最后一个，所以不需要减去1
            temp = str(excel_table_1.cell(i,j).value)
            if('\n' in temp and '\r' in temp):
                list_val[i][j] = temp.replace('\r','')
            elif ('\r' in temp):
                list_val[i][j] = temp.replace('\r','\n')
            else:
                list_val[i][j] = temp   
    dict_col ={"产品线":productline_col,
                "项目名称":proname_col,
                "本周进展":thisweekprocess_col,
                "风险":riskandprob_col,
                "状态":status_col,
                "汇报":report_col,
                "项目经理":pm_name_clo,
                "调整后计划":adustplan_col,
                "进度偏差":deviation_col,
                "核心需求":corerequire_col,
                "下周计划":nextweekplan_col,
                "初始计划":initplan_col
                }
    return [dict_col,list_val]


def write_product_head_docx(productline,file_docx):

    #处理行业产品线数据，由于段落选择没有合适的命令，所以区分行业，渠道，海外进行处理
    p_hy=file_docx.add_paragraph('')
    p_hy1 = p_hy.add_run(productline+'产品线')
    p_hy1.font.bold = True
    p_hy1.font.underline = True
    #颜色
    p_hy1.font.color.rgb = docx.shared.RGBColor(0, 176 , 80)
    # 设置中文字体
    p_hy1.font.name = u'微软雅黑'
    p_hy1.element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
    #字号
    p_hy1.font.size = Pt(11)
    
    return 1

def write_docx(read_excel_result,productline,file_docx,pro_num):
    doct_col =  read_excel_result[0]
    list_val = read_excel_result[1]
    
    pro_rep_num = pro_num

    for i in range(1,len(list_val)): 
        if( (productline in list_val[i][doct_col["产品线"]]) and list_val[i][doct_col["状态"]] == "开发中" and list_val[i][doct_col["汇报"]] == "是") :
            pro_rep_num += 1 
            temp = NumberToChinese(pro_rep_num)+"、"+ str(list_val[i][doct_col["项目名称"]]).replace("\r",'').replace('\n','')
            file_docx.add_paragraph(temp)

            temp = list_val[i][doct_col["本周进展"]]
            if(temp == ""):
                    print("error!") #考虑补充错误补充
            file_docx.add_paragraph(temp)
            p_hy = file_docx.add_paragraph("")
            temp_run = p_hy.add_run("风险以及问题：")
            temp_run.font.color.rgb = docx.shared.RGBColor(255,0,0)
            p_hy = file_docx.add_paragraph("")
            temp = list_val[i][doct_col["风险"]]
            if(temp == ""):
                temp = "1、暂无;"
            temp_run = p_hy.add_run(temp)
            temp_run.font.color.rgb = docx.shared.RGBColor(255,0,0)
            p_hy.add_run("\n")
    
    return pro_rep_num  

def write_excel_firstrow(worksheet):#第一行

    dict_col ={"产品线":0,"项目经理":1,"项目名称":2,"核心需求":3,"本周进展":4,"下周计划":5,"风险":6,"初始计划":7,"调整后计划":8,"进度偏差":9}

    worksheet.row(0).set_style(style_headtype[1])
    
    worksheet.write(0,dict_col["产品线"],"产品线",style_headtype[0])
    worksheet.col(dict_col["产品线"]).width = 10*256  ## 256为衡量单位，10表示10个字符宽度

    worksheet.write(0,dict_col["项目经理"],"项目经理",style_headtype[0])
    worksheet.col(dict_col["项目经理"]).width = 10*256

    worksheet.write(0,dict_col["项目名称"],"项目名称",style_headtype[0])
    worksheet.col(dict_col["项目名称"]).width = 21*256 

    worksheet.write(0,dict_col["本周进展"],"本周进展",style_headtype[0])
    worksheet.col(dict_col["本周进展"]).width = 40*256 

    worksheet.write(0,dict_col["下周计划"],"下周计划",style_headtype[0])
    worksheet.col(dict_col["下周计划"]).width = 40*256  

    worksheet.write(0,dict_col["风险"],"风险",style_headtype[0])
    worksheet.col(dict_col["风险"]).width = 40*256 

    worksheet.write(0,dict_col["初始计划"],"初始计划",style_headtype[0])
    worksheet.col(dict_col["初始计划"]).width = 30*256 

    worksheet.write(0,dict_col["调整后计划"],"调整后计划",style_headtype[0])
    worksheet.col(dict_col["调整后计划"]).width = 30*256 

    worksheet.write(0,dict_col["核心需求"],"核心需求",style_headtype[0])
    worksheet.col(dict_col["核心需求"]).width = 35*256  

    worksheet.write(0,dict_col["进度偏差"],"进度偏差说明",style_headtype[0])
    worksheet.col(dict_col["进度偏差"]).width = 40*256  

    return dict_col

def write_excel(read_excel_result,pro_num,dict_col_new,worksheet):
    dict_col =  read_excel_result[0]
    list_val = read_excel_result[1]
    
    pro_rep_num = pro_num 

    for i in range(0,len(list_val)): 
        if(list_val[i][dict_col["状态"]] == "开发中"):
            worksheet.row(pro_rep_num).set_style(style_bodytype[1])
            worksheet.write(pro_rep_num,dict_col_new["产品线"],list_val[i][dict_col["产品线"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["项目经理"],list_val[i][dict_col["项目经理"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["项目名称"],list_val[i][dict_col["项目名称"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["核心需求"],list_val[i][dict_col["核心需求"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["本周进展"],list_val[i][dict_col["本周进展"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["下周计划"],list_val[i][dict_col["下周计划"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["风险"],list_val[i][dict_col["风险"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["初始计划"],list_val[i][dict_col["初始计划"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["调整后计划"],list_val[i][dict_col["调整后计划"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["进度偏差"],list_val[i][dict_col["进度偏差"]],style_bodytype[0])
            pro_rep_num += 1 
    return pro_rep_num

def write_outsea_excel(read_excel_result,pro_num,dict_col_new,worksheet):
    dict_col =  read_excel_result[0]
    list_val = read_excel_result[1]
    
    pro_rep_num = pro_num 

    for i in range(0,len(list_val)): 
        if(list_val[i][dict_col["状态"]] == "开发中" and ("海外" in list_val[i][dict_col["产品线"]])):
            worksheet.row(pro_rep_num).set_style(style_bodytype[1])
            worksheet.write(pro_rep_num,dict_col_new["产品线"],list_val[i][dict_col["产品线"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["项目经理"],list_val[i][dict_col["项目经理"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["项目名称"],list_val[i][dict_col["项目名称"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["核心需求"],list_val[i][dict_col["核心需求"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["本周进展"],list_val[i][dict_col["本周进展"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["下周计划"],list_val[i][dict_col["下周计划"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["风险"],list_val[i][dict_col["风险"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["初始计划"],list_val[i][dict_col["初始计划"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["调整后计划"],list_val[i][dict_col["调整后计划"]],style_bodytype[0])
            worksheet.write(pro_rep_num,dict_col_new["进度偏差"],list_val[i][dict_col["进度偏差"]],style_bodytype[0])
            pro_rep_num += 1 
    return pro_rep_num

def setExcelStyle():
    
    #设置头一行格式

    #设置单元格的对齐方式
    alignment=xlwt.Alignment()
    alignment.horz=xlwt.Alignment.HORZ_LEFT
    alignment.vert=xlwt.Alignment.VERT_CENTER

    #设置边框
    borders=xlwt.Borders()
    borders.left=1
    borders.right=1
    borders.top=1
    borders.bottom=1

    #设置单元格的对齐方式
    #alignment1=xlwt.Alignment()
    #alignment1.horz=xlwt.Alignment.HORZ_LEFT
    #alignment1.vert=xlwt.Alignment.VERT_CENTER

    #font1的格式为表头 加粗 宋体 16号
    font1 = xlwt.Font()
    font1.name = '微软雅黑'
    font1.bold = True
    font1.height = 11*20 #字体大小为十六进制转为十进制 除以20 0x00104=320

    
    style_head = xlwt.XFStyle() # create the style
    style_head.font = font1
    style_head.alignment = alignment
    style_head.borders = borders

    #font2为  11号
    font2 = xlwt.Font()
    font2.name = '微软雅黑'
    font2.bold = False
    font2.height = 10*20 # 字体大小为十六进制转为十进制 除以20 0x00DC=220
    style_body = xlwt.XFStyle() # create the style
    style_body.font = font2
    style_body.alignment=alignment
    style_body.alignment.wrap = 1
    style_body.borders=borders


 
    tall_style_head = xlwt.easyxf('font:height 460;')  # *20   /20 才是像素高度，
    tall_style_body = xlwt.easyxf('font:height 3000;')  # 72pt
    
    style_headtype = [style_head,tall_style_head]
    style_bodytype = [style_body,tall_style_body]


    return style_headtype,style_bodytype



if __name__ == '__main__':

    #path1 = easygui.fileopenbox()
  
    #path1 = r"D:\gitRepo\work_value\智能组项目进展汇总 - 2019.xlsx"
    path1 = r"D:\基线\项目管理\项目进度\版本基线进度\智能组\智能组项目进展汇总 - 2019.xlsx"
    read_excel_result1 = open_worksheet(path1)
    
    time.sleep(1)
    
    #path2 = easygui.fileopenbox()
    
    #path2 = r"D:\gitRepo\work_value\海外组项目进展汇总.xlsx"
    path2 = r"D:\基线\项目管理\项目进度\版本基线进度\海外组\海外组项目进展汇总.xlsx"
    read_excel_result2 = open_worksheet(path2)
    
    time.sleep(1)

    
    #path3 = r"D:\gitRepo\work_value\通用组项目进展汇总 - 2019.xlsx"
    path3 = r"D:\基线\项目管理\项目进度\版本基线进度\通用组\通用组项目进展汇总 - 2019.xlsx"
   
   # path3 = easygui.fileopenbox()
    read_excel_result3 = open_worksheet(path3)
    
    time.sleep(1)

    file_docx=docx.Document()
    #file_docx.add_heading('行业组重点项目进展', 1)
    docx_style = file_docx.styles['Normal']
    # 设置西文字体
    docx_style.font.name = '微软雅黑'
    # 设置中文字体
    docx_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    pro_num = 0

   


    write_product_head_docx("渠道",file_docx)
    pro_num = 0
    pro_num = write_docx(read_excel_result1,"渠道",file_docx,pro_num)
    pro_num = write_docx(read_excel_result2,"渠道",file_docx,pro_num)
    pro_num = write_docx(read_excel_result3,"渠道",file_docx,pro_num)
    
    write_product_head_docx("行业",file_docx)
    pro_num = 0
    pro_num = write_docx(read_excel_result1,"行业",file_docx,pro_num)
    pro_num = write_docx(read_excel_result2,"行业",file_docx,pro_num)
    pro_num = write_docx(read_excel_result3,"行业",file_docx,pro_num)
    
    write_product_head_docx("专用",file_docx)
    pro_num = 0
    pro_num = write_docx(read_excel_result1,"专用",file_docx,pro_num)
    pro_num = write_docx(read_excel_result2,"专用",file_docx,pro_num)
    pro_num = write_docx(read_excel_result3,"专用",file_docx,pro_num)
    
    write_product_head_docx("海外",file_docx)
    pro_num = 0
    pro_num = write_docx(read_excel_result1,"海外",file_docx,pro_num)
    pro_num = write_docx(read_excel_result2,"海外",file_docx,pro_num)
    pro_num = write_docx(read_excel_result3,"海外",file_docx,pro_num)

    date_to = time.localtime()
    week_to = time.strftime("%U",date_to)
   
    workbook = xlwt.Workbook(encoding = 'utf-8')
    # 创建一个worksheet
    worksheet = workbook.add_sheet("汇总",cell_overwrite_ok=True)

    style_headtype,style_bodytype = setExcelStyle() # 设置全局变量,直接不做数字传递


    dict_newcol = write_excel_firstrow(worksheet)

    pro_num = 1

    pro_num = write_excel(read_excel_result1,pro_num,dict_newcol,worksheet)
    pro_num = write_excel(read_excel_result2,pro_num,dict_newcol,worksheet)
    pro_num = write_excel(read_excel_result3,pro_num,dict_newcol,worksheet)
    

    workbook_outsea = xlwt.Workbook(encoding = 'utf-8')
    # 创建一个worksheet
    worksheet_outsea = workbook_outsea.add_sheet("汇总",cell_overwrite_ok=True)

    dict_newcol = write_excel_firstrow(worksheet_outsea)

    pro_num = 1

    pro_num = write_outsea_excel(read_excel_result1,pro_num,dict_newcol,worksheet_outsea)
    pro_num = write_outsea_excel(read_excel_result2,pro_num,dict_newcol,worksheet_outsea)
    pro_num = write_outsea_excel(read_excel_result3,pro_num,dict_newcol,worksheet_outsea)

    str_doc = "本周重点项目汇总" +"_W"+ week_to +".docx"
    str_exc = "本周项目进展汇总" +"_W"+ week_to+".xls"
    str_outsea = "本周海外项目进展汇总" +"_W"+ week_to+".xls"

    file_docx.save(str_doc)
    workbook.save(str_exc)
    workbook_outsea.save(str_outsea)

    input("已经完成\n"+str_doc+"\n"+str_exc+"\n"+str_outsea)